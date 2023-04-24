import sqlite3
from docx import Document
#from openpyxl import *
from datetime import datetime
from datetime import date
from shutil import copyfile
from docx import *
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn

con = sqlite3.connect('Certificados.db')
cur=con.cursor()
cur.execute('PRAGMA foreign_keys = ON')

cont = int()

def criarQualidade():
    sql="""CREATE TABLE IF NOT EXISTS qualidade (
           NF text NOT NULL,
           bitola text NOT NULL PRIMARY KEY,
           data text NOT NULL,
           lote text NOT NULL,
           empresa text NOT NULL,
           estado text NOT NULL,
           cnpj text NOT NULL,
           bitola1 text NOT NULL
           );"""

    try:
        cur = con.cursor()
        cur.execute(sql)
    except Error as e:
        print(e)

    return


def criarCliente():
    sql = """CREATE TABLE IF NOT EXISTS clientes (
             cnpj text NOT NULL,
             ie text NOT NULL,
             codigo text NOT NULL,
             nome text NOT NULL,
             tipolog text NOT NULL,
             logradouro text NOT NULL,
             numero text,
             complemento text,
             bairro text,
             estado text,
             cidade text
             );"""

    try:
        cur = con.cursor()
        cur.execute(sql)
    except Error as e:
        print(e)

    return

def criarItens():
    sql = """CREATE TABLE IF NOT EXISTS itens (
             codigo text NOT NULL,
             codigo_alt text,
             desc text NOT NULL,
             unidade text NOT NULL             
             );"""

    try:
        cur = con.cursor()
        cur.execute(sql)
    except Error as e:
        print(e)

    return

def criarContador():
    sql = """CREATE TABLE IF NOT EXISTS contador (
             versao integer NOT NULL,
             id integer NOT NULL
             );"""

    try:
        cur = con.cursor()
        cur.execute(sql)
    except Error as e:
        print(e)

    return

def criarIMGQualidade():
    sql = """CREATE TABLE IF NOT EXISTS img (
             bitola text NOT NULL,
             certificado BLOB
             );"""

    try:
        cur = con.cursor()
        cur.execute(sql)
    except Error as e:
        print(e)

    return
    

def cadastrarVersao(versao,id):
    sql = "INSERT OR IGNORE INTO contador(versao,id)VALUES(?,?)"
    dados = (versao,id, )
    if versao == '':
        raise
    else:
        cur = con.cursor()
        cur.execute(sql,dados)
        con.commit()

    return

def cadastrarCliente(cnpj,ie,codigo,nome,tipolog,logradouro,numero,complemento,bairro,estado,cidade):
    sql = "INSERT INTO clientes(cnpj,ie,codigo,nome,tipolog,logradouro,numero,complemento,bairro,estado,cidade)VALUES(?,?,?,?,?,?,?,?,?,?,?)"
    dados = (cnpj,ie,codigo,nome,tipolog,logradouro,numero,complemento,bairro,estado,cidade)
    if cnpj == '' or ie == '' or codigo == '' or nome == '':
        raise
    if numero == '':
        numero = "-"
    if complemento == '':
        complemento = "-"
    if tipolog == '':
        tipolog = "-"
    if logradouro == '':
        logradouro = "-"
    if bairro == '':
        bairro = "-"
    if estado == '':
        estado = "-"
    if cidade == '':
        cidade = "-"
    else:
        cur = con.cursor()
        cur.execute(sql,dados)
        con.commit()

    return

#cur.execute("DROP TABLE IF EXISTS itens")
#criarItens()



def cadastrarItens(codigo,codigo_alt,desc,unidade):
    sql = "INSERT INTO itens(codigo,codigo_alt,desc,unidade)VALUES(?,?,?,?)"
    dados = (codigo,codigo_alt,desc,unidade)
    if codigo == '' or desc == "" or unidade == "":
        raise
    if str(codigo_alt) == str(codigo):
        codigo_alt = "-"
    else:
        cur = con.cursor()
        cur.execute(sql,dados)
        con.commit()

    return


def importarExcel():
    #cur.execute("DROP TABLE IF EXISTS CAMA")
    excel = load_workbook('ITENS.xlsx')
    plan1=excel['ITENS']

    for i in plan1:
        l1=i[0].value
        l2=i[1].value
        l3=i[2].value
        l4=i[3].value
        
        print(i)
        cadastrarItens(l1,l2,l3,l4)

#importarExcel()

def cadastrarQualidade(nf,bitola,data,lote,empresa,estado,cnpj,bitola1):
    sql = "INSERT INTO qualidade(NF,bitola,data,lote,empresa,estado,cnpj,bitola1)VALUES(?,?,?,?,?,?,?,?)"
    dados = (nf,bitola,data,lote,empresa,estado,cnpj,bitola1)
    if nf == "" or bitola == "" or data == "" or lote == "" or empresa== "" or estado == "" or cnpj == "":
        raise
    else:
        cur = con.cursor()
        cur.execute(sql,dados)
        con.commit()

    return


def cadastrarIMGQualidade(bitola,certificado):
    sql = "INSERT INTO img(bitola,certificado)VALUES(?,?)"
    
    if certificado != '':
        certificado1 = convertToBinaryData(certificado)
    else:
        certificado1 = certificado

    dados = (bitola,certificado1)
    cur = con.cursor()
    cur.execute(sql,dados)
    con.commit()

    return

def convertToBinaryData(filename):
    with open(filename,'rb') as file:
        blobData = file.read()
    return blobData


def writeTofile(data,filename):
    with open(filename, 'wb') as file:
        file.write(data)

#cadastrarIMGQualidade('M14',r'C:\Users\Calebe\Downloads\CERTIFICADO - SIMEC - 9.16.jpeg')

###################################################################################################################################################################################################################################################################################################################################################################################################################################################################################################################################################################


def pegarCliente(cnpj):
    cur = con.cursor()
    cur.execute("SELECT * FROM clientes WHERE cnpj=?",(cnpj, ))
    r1 = cur.fetchall()
    cliente = r1[0][3]
    end = r1[0][4] + " " + r1[0][7] + ", " + r1[0][8] + " " + r1[0][9]
    cidadeest = r1[0][6] + " / " + r1[0][5]
    cnpj = r1[0][0]
    

    return cliente, end, cidadeest, cnpj


def quantosItens(QTD):

    QTDITENS = QTD

    return QTDITENS


def pegarItem(codigo, codigo_alt):
    r2 = []
    cur = con.cursor()
    cur.execute("SELECT * FROM itens WHERE codigo=? or codigo_alt=?",(codigo, codigo_alt, ))
    r1 = cur.fetchall()
    r2.append(r1[0][2])
    r2.append(r1[0][3])
    return r2

def gerarLis(lis):
    for i in range(len(lis)):
        if "3/16" in lis[i][1]:
            lis[i].insert(2,'3/16')
            lis[i].append('(ROSCA MAQUIA DIREITA 24 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 120 KGF')
        if "1/4" in lis[i][1] and "1 1/4" not in lis[i][1]:
            lis[i].insert(2,'1/4')
            lis[i].append('(ROSCA MAQUIA DIREITA 20 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 220 KGF')
        if "5/16" in lis[i][1]:
            lis[i].insert(2,'5/16')
            lis[i].append('(ROSCA MAQUIA DIREITA 18 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 380 KGF')
        if "3/8" in lis[i][1]:
            lis[i].insert(2,'3/8')
            lis[i].append('(ROSCA MAQUIA DIREITA 16 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 540 KGF')
        if "7/16" in lis[i][1]:
            lis[i].insert(2,'7/16')
            lis[i].append('(ROSCA MAQUIA DIREITA 14 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 700 KGF')
        if "1/2" in lis[i][1] and "NC" in lis[i][1]:
            lis[i].insert(2,'1/2')
            lis[i].append('(ROSCA MAQUIA DIREITA 13 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 900 KGF')
        if "1/2" in lis[i][1] and "WW" in lis[i][1]:
            lis[i].insert(2,'1/2')
            lis[i].append('(ROSCA MAQUIA DIREITA 12 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 900 KGF')
        if "5/8" in lis[i][1]:
            lis[i].insert(2,'5/8')
            lis[i].append('(ROSCA MAQUIA DIREITA 11 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 1500 KGF')
        if "3/4" in lis[i][1]:
            lis[i].insert(2,'3/4')
            lis[i].append('(ROSCA MAQUIA DIREITA 10 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 2200 KGF')
        if "7/8" in lis[i][1]:
            lis[i].insert(2,'7/8')
            lis[i].append('(ROSCA MAQUIA DIREITA 9 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 3000 KGF')
        if "1 X" in lis[i][1]:
            lis[i].insert(2,'1')
            lis[i].append('(ROSCA MAQUIA DIREITA 8 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 4000 KGF')
        if "1 1/8" in lis[i][1]:
            lis[i].insert(2,'1 1/8')
            lis[i].append('(ROSCA MAQUIA DIREITA 7 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 5000 KGF')
        if "1 1/4" in lis[i][1]:
            lis[i].insert(2,'1 1/4')
            lis[i].append('(ROSCA MAQUIA DIREITA 7 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 6000 KGF')
        if "1 1/2" in lis[i][1]:
            lis[i].insert(2,'1 1/2')
            lis[i].append('(ROSCA MAQUIA DIREITA 6 F.P.P.)')
            lis[i].append('CARGA DE TRABALHO: 9000 KGF')
        if "M6" in lis[i][1]:
            lis[i].insert(2,'M6')
            lis[i].append('(PASSO ROSCA 1,00)')
            lis[i].append('CARGA DE TRABALHO: 200 KGF')
        if "M8" in lis[i][1]:
            lis[i].insert(2,'M8')
            lis[i].append('(PASSO ROSCA 1,25)')
            lis[i].append('CARGA DE TRABALHO: 380 KGF')
        if "M10" in lis[i][1]:
            lis[i].insert(2,'M10')
            lis[i].append('(PASSO ROSCA 1,50)')
            lis[i].append('CARGA DE TRABALHO: 560 KGF')
        if "M12" in lis[i][1]:
            lis[i].insert(2,'M12')
            lis[i].append('(PASSO ROSCA 1,75)')
            lis[i].append('CARGA DE TRABALHO: 860 KGF')
        if "M14" in lis[i][1]:
            lis[i].insert(2,'M14')
            lis[i].append('(PASSO ROSCA 2,00)')
            lis[i].append('CARGA DE TRABALHO: 1150 KGF')
        if "M16" in lis[i][1]:
            lis[i].insert(2,'M16')
            lis[i].append('(PASSO ROSCA 2,00)')
            lis[i].append('CARGA DE TRABALHO: 1650 KGF')
        if "M20" in lis[i][1]:
            lis[i].insert(2,'M20')
            lis[i].append('(PASSO ROSCA 2,50)')
            lis[i].append('CARGA DE TRABALHO: 2500 KGF')
        if "M24" in lis[i][1]:
            lis[i].insert(2,'M24')
            lis[i].append('(PASSO ROSCA 3,00)')
            lis[i].append('CARGA DE TRABALHO: 3500 KGF')
        if "M30" in lis[i][1]:
            lis[i].insert(2,'M30')
            lis[i].append('(PASSO ROSCA 3,50)')
            lis[i].append('CARGA DE TRABALHO: 5500 KGF')
    return lis


#for i in range(QTDITENS):
#    lis.append([input("Digite a quantidade: "),input("Digite o tipo: "), input("Digite a Bitola: "), input("Digite o comprimento: "), input("Digite o banho: "), f'{i+1}'])


doc = Document()



"_________________________________________________Fonte e Margem____________________________________________________"

def fontMarg(doc):
    font = doc.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(11)


    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(0.75)
        section.left_margin = Cm(1.75)
        section.right_margin = Cm(1.5)
        return

"__________________________________________________________________________________________________________________"
"________________________________________________Cabeçalho e Logo_______________________________________________"

def cabeçalhoLogo(doc):
    cur = con.cursor()
    cur.execute("SELECT versao FROM contador")
    r1=cur.fetchall()
    section = doc.sections[0]
    headers = section.header
    section.header_distance = Cm(1)
    paragraph = headers.paragraphs[0]

    b = paragraph.add_run()
    b.add_picture("logoMUBEC.png", width=Inches(3.5))

    c = paragraph.add_run()
    c.text = '\t' + f"   ERTIFICADO DE CONFORMIDADE N° {r1[0][0]}" # For center align of text
    font = c.font
    font.name = 'Arial'
    font.size = Pt(12.5)
    return

"__________________________________________________________________________________________________________________"


"___________________________________________Dados Cliente____________________________________________________"

def dadosCliente(doc,cliente,end,cidadeest,cnpj,NF,data):
    a = doc.add_paragraph()
    a.paragraph_format.space_after = Pt(3)
    a = doc.add_paragraph()
    a.add_run(f'Certificado referente à nota fiscal: Nº {NF} de {data}').bold = True


    a = doc.add_paragraph(f"Cliente: {cliente}")
    a.paragraph_format.space_after = Pt(3)
    a=doc.add_paragraph(f"               {end}")
    a.paragraph_format.space_after = Pt(3)
    a=doc.add_paragraph(f"               {cidadeest} ")
    a.paragraph_format.space_after = Pt(3)
    a=doc.add_paragraph(f"CNPJ: {cnpj}")
    a.paragraph_format.space_after = Pt(3)
    a=doc.add_paragraph("Pedido: -")
    a.paragraph_format.space_after = Pt(3)
    a=doc.add_paragraph("Descrição do Produto")
    a.paragraph_format.space_after = Pt(3)
    return


"__________________________________________________________________________________________________________________"

"_______________________________________________Criação da tabela__________________________________________________"


def criarTabela(doc,QTDITENS):
    table = doc.add_table(rows=QTDITENS+1, cols=4)


    row = table.columns[0].cells
    row[0].text = 'ITEM'
    row[0].width = Cm(1.5)
    row[0].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER


    row = table.columns[1].cells
    row[0].text = 'QUANT.'
    row[0].width = Cm(1.5)
    row[0].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table.columns[2].cells
    row[0].text = 'UNID'
    row[0].width = Cm(1.5)
    row[0].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table.columns[3].cells
    row[0].text = 'DESCRIÇÃO'
    row[0].width = Cm(15)
    row[0].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table.rows[0]
    row.height = Cm(2)


    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    font = run.font
                    font.size= Pt(8.5)
                    font.name = 'Cambria'
    return table
"__________________________________________________________________________________________________________________"

"_________________________________________Inserção dos itens_______________________________________________________"

def inserirItens(doc,lis,table):
    for i,z in zip(range(1, len(table.rows)),range(len(lis))):
        for j in range(len(table.columns)):
            row = table.columns[j].cells
            if j == 0:
                row[i].text = lis[z][4]
                row[i].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j == 1:
                row[i].text = lis[z][0]
                row[i].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j == 2:
                row[i].text = lis[z][3]
                row[i].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j == 3:
                row[i].text = lis[z][1] + " " + lis[z][5]
                a = row[i].add_paragraph(lis[z][6])
                row[i].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                a.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(1, len(table.rows)):
        row = table.rows[i]
        row.height = Cm(1)


    for i in range(1, len(table.rows)):
        for cell in table.rows[i].cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.font.bold = False
                    font = run.font
                    font.size= Pt(8)
                    font.name = 'Arial'

    for col in table.columns:
        for cell in col.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.style = 'Table Grid'
    return

"__________________________________________________________________________________________________________________"


"__________________________________________________________________________________________________________________"

def inserirMUBEC(doc):
    a = doc.add_paragraph("Fabricante: ")
    a.paragraph_format.space_before = Pt(15)
    a.paragraph_format.space_after = Pt(3)
    a.add_run('MUBEC INDÚSTRIA E COMÉRCIO LTDA').bold = True


    a=doc.add_paragraph("CNPJ: ")
    a.paragraph_format.space_before = Pt(5)
    a.paragraph_format.space_after = Pt(3)
    a.add_run('00.604.905/0001-70').bold = True

    doc.add_paragraph()
    return
"__________________________________________________________________________________________________________________"

"______________________________________________Separação de itens por bitola_______________________________________"

def separarBitolas(lis,QTDITENS):

    itens = {}
    chaves = []
    for i in range(len(lis)):
        if lis[i][2] not in chaves:
            itens[lis[i][2]] = []
            chaves = [*itens]
        for j in range(len(lis)):
            if lis[i][2] == lis[j][2] and lis[i][4] != lis[j][4]:
                if lis[i][4] not in itens[lis[i][2]]:
                    itens[lis[i][2]].append(lis[i][4])
            elif lis[i][2] != lis [j][2] and lis[i][4] not in itens[lis[i][2]]:
                itens[lis[i][2]].append(lis[i][4])
            elif QTDITENS == 1:
                itens[lis[i][2]].append(lis[i][4])
    return itens,chaves


doc.add_paragraph()
"__________________________________________________________________________________________________________________"


"___________________________________________________Certificado por itens________________________________________________"

def inserirCertificado(doc,itens,chaves):
    impedidor = []
    for i in chaves:
        r = str()
        for j in itens[i]:
            if len(itens[i]) > 1 and i not in impedidor:
                impedidor.append(i)
                for x in itens[i]:
                    if r == '':
                        r = x
                    else:
                        r = r + ', ' + x
                a = doc.add_paragraph()
                b = a.add_run(f"Itens {r}")
                b.bold = True
                a.paragraph_format.space_before = Pt(0)
                a.paragraph_format.space_after = Pt(0)
                font = b.font
                font.name = 'Arial'
                font.size = Pt(10)
                
                itens[i][:-1]

                for h in itens.keys():
                    if h == i:
                        keys = i

                cur = con.cursor()
                cur.execute("SELECT NF,data,lote,empresa,estado,cnpj,bitola1 FROM qualidade WHERE bitola=?",(keys, ))
                r1 = cur.fetchall()
                r2=[]
                for l in r1:
                    r2.append(l)                    

                a = doc.add_paragraph()
                a.paragraph_format.space_before = Pt(0)
                a.paragraph_format.space_after = Pt(0)
                b = a.add_run(f"Nº Certificado de origem: {keys} – NF {r2[0][0]} de {r2[0][1]} – Corrida {r2[0][2]} ")
                font = b.font
                font.name = 'Arial'
                font.size = Pt(10)
                
                a = doc.add_paragraph()
                b = a.add_run(f"{r2[0][3]} / {r2[0][4]}  -  CNPJ: {r2[0][5]}")
                a.paragraph_format.space_before = Pt(0)
                a.paragraph_format.space_after = Pt(0)
                font = b.font
                font.name = 'Arial'
                font.size = Pt(10)
                
                doc.add_paragraph()

                cur.execute("SELECT certificado FROM img WHERE bitola = ?", (keys, ))
                r3=cur.fetchall()
                photoPath = r"MP\\" + r2[0][6] + ".PNG"
                writeTofile(r3[0][0],photoPath)

                
            elif len(itens[i]) == 1 and i not in impedidor:
                impedidor.append(i)
                for x in itens[i]:
                    if r == '':
                        r = x
                    else:
                        r = r + ', ' + x
                a = doc.add_paragraph()
                b = a.add_run(f"Itens {r}")
                b.bold = True
                a.paragraph_format.space_before = Pt(0)
                a.paragraph_format.space_after = Pt(0)
                font = b.font
                font.name = 'Arial'
                font.size = Pt(10)
                
                #itens[i][:-1]

                for h in itens.keys():
                    if h == i:
                        keys = i
                cur = con.cursor()
                cur.execute("SELECT NF,data,lote,empresa,estado,cnpj,bitola1 FROM qualidade WHERE bitola=?",(keys, ))
                r1 = cur.fetchall()
                r2=[]
                for l in r1:
                    r2.append(l)
                    
                a = doc.add_paragraph()
                b = a.add_run(f"Nº Certificado de origem: {keys} – NF {r2[0][0]} de {r2[0][1]} – Corrida {r2[0][2]} ")
                a.paragraph_format.space_before = Pt(0)
                a.paragraph_format.space_after = Pt(0)
                font = b.font
                font.name = 'Arial'
                font.size = Pt(10)
                
                
                a = doc.add_paragraph()
                b = a.add_run(f"{r2[0][3]} / {r2[0][4]}  -  CNPJ: {r2[0][5]}")
                a.paragraph_format.space_before = Pt(0)
                a.paragraph_format.space_after = Pt(0)
                font = b.font
                font.name = 'Arial'
                font.size = Pt(10)
                
                doc.add_paragraph()


                cur.execute("SELECT certificado FROM img WHERE bitola = ?", (keys, ))
                r3=cur.fetchall()
                photoPath = r"MP\\" + r2[0][6] + ".PNG"
                writeTofile(r3[0][0],photoPath)
    return

def criarBanho(doc, lis):
    cont1=0
    cont=0
    for i in range(len(lis)):
        if "GE" in lis[i][1] and cont1==0:
            a = doc.add_paragraph("Galvanização: ")
            a.add_run("JJ LESTE GALVANIZAÇÃO LTDA").bold = True
            a.paragraph_format.space_before = Pt(10)
            a.paragraph_format.space_after = Pt(0)
            a=doc.add_paragraph()
            a.add_run("RUA PARTICULAR TIMÃO, 76- CID TIRADENTES").bold = True
            a.paragraph_format.space_before = Pt(0)
            a.paragraph_format.space_after = Pt(0)
            a = doc.add_paragraph("CNPJ: ")
            a.add_run("26.412.069/0001-16"). bold = True
            a.paragraph_format.space_before = Pt(0)
            a.paragraph_format.space_after = Pt(0)
            a = doc.add_paragraph("Passivação: ")
            a.add_run("Azul"). bold = True
            a.paragraph_format.space_before = Pt(0)
            a.paragraph_format.space_after = Pt(0)
            a = doc.add_paragraph("Camada: ")
            a.add_run("8 Microons"). bold = True
            a.paragraph_format.space_before = Pt(0)
            a.paragraph_format.space_after = Pt(10)
            cont1 += 1
            
        if "GF" in lis[i][1] and cont == 0:
            a = doc.add_paragraph("Galvanização: ")
            a.add_run("GALVALLE INDUSTRIA E COMERCIO LTDA").bold = True
            a.paragraph_format.space_before = Pt(10)
            a.paragraph_format.space_after = Pt(0)
            a=doc.add_paragraph()
            a.add_run("RUA LANDRI SALES, 1633- CID ARACILIA").bold = True
            a.paragraph_format.space_before = Pt(0)
            a.paragraph_format.space_after = Pt(0)
            a = doc.add_paragraph("CNPJ: ")
            a.add_run("12.882.845/0001-37"). bold = True
            a.paragraph_format.space_before = Pt(0)
            a.paragraph_format.space_after = Pt(0)
            a = doc.add_paragraph("Passivação: ")
            a.add_run("Galvanizado à fogo"). bold = True
            a.paragraph_format.space_before = Pt(0)
            a = doc.add_paragraph("NBR 6323")
            a.paragraph_format.space_after = Pt(10)
            cont +=1
            


"__________________________________________________________________________________________________________________"       

def finalmentes(doc):

    a = doc.add_paragraph()
    b = a.add_run("Certificamos o envio do produto acima, através da Nota Fiscal em referência. Material Produzido e Inspecionado de acordo com todas exigências.")
    a.paragraph_format.space_before = Pt(0)
    a.paragraph_format.space_after = Pt(0)
    font = b.font
    font.name = 'Arial'
    font.size = Pt(10)

    a = doc.add_paragraph()
    a.paragraph_format.space_before = Pt(0)
    a.paragraph_format.space_after = Pt(0)
    a = doc.add_paragraph()
    a.paragraph_format.space_before = Pt(0)
    a.paragraph_format.space_after = Pt(0)


    a = doc.add_paragraph()
    b = a.add_run('Gerado eletronicamente, dispensa assinatura.')
    b.underline = True
    a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font = b.font
    font.name = 'Arial'
    font.size = Pt(12)
    a.paragraph_format.space_before = Pt(0)
    a.paragraph_format.space_after = Pt(0)


    a = doc.add_paragraph()
    b = a.add_run('QUALIDADE: Calebe Louzada')
    a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font = b.font
    font.name = 'Arial'
    font.size = Pt(11.5)
    a.paragraph_format.space_before = Pt(0)
    a.paragraph_format.space_after = Pt(0)
    return

"__________________________________________________________________________________________________________________"
def rodape(doc):
    section = doc.sections[0]
    footers = section.footer
    section.footer_distance = Cm(0.8)
    a = footers.paragraphs[0]
    b = a.add_run("Av. do Estado , 6894 – Cambuci – São Paulo  / SP   Fone : (11) 2271-2900  /   e-mail  :  qualidade@mubec.com.br")
    a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    font = b.font
    font.name = 'Calibri'
    font.size = Pt(11)
    a.paragraph_format.space_before = Pt(0)
    a.paragraph_format.space_after = Pt(0)
    return
"__________________________________________________________________________________________________________________"


"_______________________________________________________Margem____________________________________________________"

def margem(doc):
    sec_pr = doc.sections[0]._sectPr
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single') # a single line
        border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), 'auto')
        pg_borders.append(border_el) # register single border to border el
    sec_pr.append(pg_borders) # apply border changes to section
"__________________________________________________________________________________________________________________"

def rodar(cliente,end,cidadeest,cnpj,QTDITENS,lis,NF,data):
    cur = con.cursor()
    cur.execute("SELECT versao FROM contador")
    r1=cur.fetchall()
    num = r1[0][0] + 1

    
    doc = Document()
    fontMarg(doc)
    cabeçalhoLogo(doc)
    dadosCliente(doc,cliente,end,cidadeest,cnpj,NF,data)
    tab1 = criarTabela(doc,QTDITENS)
    inserirItens(doc,lis,tab1)
    inserirMUBEC(doc)
    itens, chaves = separarBitolas(lis,QTDITENS)
    inserirCertificado(doc,itens,chaves)
    criarBanho(doc,lis)
    finalmentes(doc)
    rodape(doc)
    margem(doc)
    cur.execute("""UPDATE contador
                   SET versao = ?
                   WHERE id = 1""",(num, ))
    con.commit()

    doc.save(rf'Certificado\\Certificado de Qualidade N°{r1[0][0]}.docx')
    os.system(f'Certificado de Qualidade N°{r1[0][0]}.docx')

def abrirDoc():
    cur = con.cursor()
    cur.execute("SELECT versao FROM contador")
    r1=cur.fetchall()
    i = int(r1[0][0])-1

    os.system(f'Certificado de Qualidade N°{i}.docx')

